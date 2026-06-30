# ==========================================================
# AI SMART DOCUMENT ANALYZER
# docx_reader.py
# ==========================================================

from docx import Document


def extract_docx_text(uploaded_file):
    """
    Extract text from DOCX.
    """

    try:

        document = Document(uploaded_file)

        text = ""

        for paragraph in document.paragraphs:

            text += paragraph.text + "\n"

        return text

    except Exception as e:

        return f"Error reading DOCX: {e}"


def get_docx_info(uploaded_file):
    """
    Return document information.
    """

    try:

        document = Document(uploaded_file)

        return {
            "paragraphs": len(document.paragraphs)
        }

    except Exception:

        return {
            "paragraphs": 0
        }